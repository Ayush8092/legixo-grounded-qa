"""Multi-format text extraction for ingestion.

Every supported format is normalized down to the same thing: a single plain
text string, handed to `app.chunking`'s existing Markdown-aware parser.
There is exactly one chunking/embedding/upsert pipeline (see
`app/chunking.py`, `app/ingestion.py`) — this module's only job is "get
plain text out of the file", not "understand documents":

- `.md`  — read as UTF-8 text. Keeps using `#` / `##` headings for
  structured, per-section chunking (unchanged from before this change).
- `.txt` — read as UTF-8 text. Plain text has no `#`/`##` headings, so it
  falls into the same single "header" chunk fallback that `chunking.py`
  already uses for a `.md` file with no `##` sections (see
  `04_statute_style_excerpt_fictional.md` in the sample corpus, which has
  no `##` headings and is chunked the same way).
- `.pdf`  — text extracted with `pypdf` (pure-Python, no native/system deps).
- `.docx` — text extracted with `python-docx` (pure-Python, no native/system
  deps), including any table cell text.

`.pdf` and `.docx` extraction also produces plain text with no markdown
structure, so they land on the same "header" chunk fallback as `.txt`. This
is what keeps ingestion a single pipeline instead of four separate ones:
format-specific code lives only in this module, and only up to the point of
producing a plain string.
"""

from pathlib import Path

SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}

# Corpus metadata files that are documentation, not knowledge to embed.
# Kept as an explicit allowlist-adjacent exclusion (rather than relying on
# extension alone) now that .txt is a supported ingestable format — without
# this, README.txt would silently start being embedded as a document.
EXCLUDED_FILENAMES = {"readme.txt", "readme.md"}


class LoaderError(ValueError):
    """Raised for an unsupported extension, a file that can't be read or
    parsed, or one that extracts to no usable text.

    Ingestion must fail loudly and clearly on these — never silently
    produce an empty chunk, and never silently skip a file the caller
    expected to be ingested.
    """


def _load_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise LoaderError(f"{path.name}: not valid UTF-8 text") from exc
    except OSError as exc:
        raise LoaderError(f"{path.name}: could not be read ({exc})") from exc


def _load_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - depends on install state
        raise LoaderError(
            "PDF ingestion requires the 'pypdf' package — "
            "see requirements.txt (pip install -r requirements.txt)"
        ) from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise LoaderError(f"{path.name}: could not be read as a PDF ({exc})") from exc

    pages: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:
            raise LoaderError(
                f"{path.name}: failed extracting text from page {page_number} ({exc})"
            ) from exc
    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - depends on install state
        raise LoaderError(
            "DOCX ingestion requires the 'python-docx' package — "
            "see requirements.txt (pip install -r requirements.txt)"
        ) from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise LoaderError(f"{path.name}: could not be read as a DOCX ({exc})") from exc

    parts: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


_LOADERS = {
    ".md": _load_text_file,
    ".txt": _load_text_file,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
}


def load_text(path: Path) -> str:
    """Extract plain text from a supported corpus file (`.md`/`.txt`/`.pdf`/`.docx`).

    Raises `LoaderError` for an unsupported extension, a file that can't be
    read/parsed, or one that yields no usable text after extraction — never
    returns an empty string, so callers never have to special-case "did
    this silently fail".
    """
    suffix = path.suffix.lower()
    loader = _LOADERS.get(suffix)
    if loader is None:
        raise LoaderError(
            f"{path.name}: unsupported file type '{suffix}' "
            f"(supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))})"
        )

    text = loader(path)

    if not text or not text.strip():
        raise LoaderError(
            f"{path.name}: extracted no text (empty or unreadable document)"
        )
    return text
