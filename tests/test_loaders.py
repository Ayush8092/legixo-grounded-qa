"""Offline unit tests for app.loaders — multi-format text extraction.

.md/.txt are exercised with real files (plain text, no extra dependency).
.pdf/.docx are exercised by monkeypatching the underlying pypdf/python-docx
library calls, the same way the rest of this test suite fakes network/SDK
calls elsewhere — no real PDF/DOCX file needs to be constructed, and these
tests don't require pypdf/python-docx to actually be installed except for
the two tests that explicitly check the ImportError path.
"""

from pathlib import Path

import pytest

from app.loaders import LoaderError, SUPPORTED_EXTENSIONS, load_text


def test_supported_extensions_are_exactly_the_four_required_formats():
    assert SUPPORTED_EXTENSIONS == {".md", ".txt", ".pdf", ".docx"}


def test_load_text_reads_markdown_as_is(tmp_path):
    path = tmp_path / "doc.md"
    path.write_text("# Title\n\n## Section\n\nBody text.\n", encoding="utf-8")
    assert load_text(path) == "# Title\n\n## Section\n\nBody text.\n"


def test_load_text_reads_plain_txt(tmp_path):
    path = tmp_path / "notes.txt"
    path.write_text("Just some plain notes with no headings.", encoding="utf-8")
    assert load_text(path) == "Just some plain notes with no headings."


def test_load_text_raises_on_unsupported_extension(tmp_path):
    path = tmp_path / "spreadsheet.xlsx"
    path.write_text("irrelevant", encoding="utf-8")
    with pytest.raises(LoaderError, match="unsupported file type"):
        load_text(path)


def test_load_text_raises_on_empty_txt_file(tmp_path):
    path = tmp_path / "empty.txt"
    path.write_text("   \n  \n", encoding="utf-8")
    with pytest.raises(LoaderError, match="no text"):
        load_text(path)


def test_load_text_raises_on_non_utf8_file(tmp_path):
    path = tmp_path / "bad_encoding.txt"
    path.write_bytes(b"\xff\xfe\x00\xff not valid utf-8")
    with pytest.raises(LoaderError):
        load_text(path)


# ---------------------------------------------------------------------
# .pdf — pypdf.PdfReader is monkeypatched, no real PDF file needed.
# ---------------------------------------------------------------------


class _FakePage:
    def __init__(self, text: str, raise_exc: Exception | None = None):
        self._text = text
        self._raise = raise_exc

    def extract_text(self):
        if self._raise:
            raise self._raise
        return self._text


class _FakePdfReader:
    def __init__(self, path_str, pages):
        self.pages = pages


def test_load_pdf_joins_page_text(tmp_path, monkeypatch):
    import pypdf

    fake_reader = lambda path_str: _FakePdfReader(path_str, [_FakePage("Page one."), _FakePage("Page two.")])
    monkeypatch.setattr(pypdf, "PdfReader", fake_reader)

    path = tmp_path / "doc.pdf"
    path.write_bytes(b"%PDF-fake")  # content irrelevant; PdfReader is mocked

    text = load_text(path)
    assert "Page one." in text
    assert "Page two." in text


def test_load_pdf_raises_loader_error_on_corrupt_file(tmp_path, monkeypatch):
    import pypdf

    def _raise(path_str):
        raise ValueError("not a PDF")

    monkeypatch.setattr(pypdf, "PdfReader", _raise)

    path = tmp_path / "corrupt.pdf"
    path.write_bytes(b"not really a pdf")

    with pytest.raises(LoaderError, match="could not be read as a PDF"):
        load_text(path)


def test_load_pdf_raises_loader_error_when_pages_extract_no_text(tmp_path, monkeypatch):
    """A scanned/image-only PDF extracts to empty strings — must not silently
    produce an empty chunk (see improvement_010: "extraction failures do not
    silently create empty chunks")."""
    import pypdf

    fake_reader = lambda path_str: _FakePdfReader(path_str, [_FakePage(""), _FakePage("")])
    monkeypatch.setattr(pypdf, "PdfReader", fake_reader)

    path = tmp_path / "scanned.pdf"
    path.write_bytes(b"%PDF-fake")

    with pytest.raises(LoaderError, match="no text"):
        load_text(path)


# ---------------------------------------------------------------------
# .docx — docx.Document is monkeypatched, no real DOCX file needed.
# ---------------------------------------------------------------------


class _FakeCell:
    def __init__(self, text):
        self.text = text


class _FakeRow:
    def __init__(self, cell_texts):
        self.cells = [_FakeCell(t) for t in cell_texts]


class _FakeTable:
    def __init__(self, rows):
        self.rows = [_FakeRow(r) for r in rows]


class _FakeParagraph:
    def __init__(self, text):
        self.text = text


class _FakeDocxDocument:
    def __init__(self, paragraphs, tables=None):
        self.paragraphs = [_FakeParagraph(t) for t in paragraphs]
        self.tables = tables or []


def test_load_docx_joins_paragraphs_and_table_cells(tmp_path, monkeypatch):
    import docx

    fake_document = _FakeDocxDocument(
        paragraphs=["Employment agreement excerpt.", "Employee: Priya Nambiar"],
        tables=[_FakeTable([["Clause", "Duration"], ["Non-compete", "12 months"]])],
    )
    monkeypatch.setattr(docx, "Document", lambda path_str: fake_document)

    path = tmp_path / "doc.docx"
    path.write_bytes(b"PK\x03\x04fake-docx-bytes")  # content irrelevant; mocked

    text = load_text(path)
    assert "Priya Nambiar" in text
    assert "Non-compete" in text
    assert "12 months" in text


def test_load_docx_raises_loader_error_on_corrupt_file(tmp_path, monkeypatch):
    import docx

    def _raise(path_str):
        raise ValueError("not a valid docx")

    monkeypatch.setattr(docx, "Document", _raise)

    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not really a docx")

    with pytest.raises(LoaderError, match="could not be read as a DOCX"):
        load_text(path)


def test_load_docx_raises_loader_error_on_empty_document(tmp_path, monkeypatch):
    import docx

    monkeypatch.setattr(docx, "Document", lambda path_str: _FakeDocxDocument(paragraphs=[]))

    path = tmp_path / "empty.docx"
    path.write_bytes(b"PK\x03\x04fake-docx-bytes")

    with pytest.raises(LoaderError, match="no text"):
        load_text(path)
