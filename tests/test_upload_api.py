"""Offline tests for POST /upload (app.main).

`ingest_corpus` is replaced with a fake so these tests need no API keys and
make no network/Pinecone/Gemini calls — only the per-file validation,
sanitization, and save-to-disk behavior of the endpoint itself is under
test here (the ingestion pipeline itself is covered by test_ingestion.py).
"""

import io

import pytest
from fastapi.testclient import TestClient

from app import main
from app.config import Settings


def _settings(tmp_path, **overrides) -> Settings:
    base = dict(
        GEMINI_API_KEY="x",
        GROQ_API_KEY="x",
        PINECONE_API_KEY="x",
        # Uploaded files must land in UPLOAD_DIR, never CORPUS_DIR — see
        # docs/architecture.md "Two knowledge bases, one logical corpus".
        # CORPUS_DIR is deliberately left at its default here (it's never
        # written to by these tests) rather than also pointed at tmp_path,
        # so a bug that accidentally reintroduces writes to corpus_dir
        # would show up as files landing somewhere unexpected instead of
        # silently passing.
        UPLOAD_DIR=str(tmp_path),
    )
    base.update(overrides)
    return Settings(**base)


class _FakeIngestSummary(dict):
    """A stand-in for ingest_corpus()'s return dict with sensible defaults."""

    def __init__(self, **overrides):
        base = dict(
            corpus_dir="ignored",
            upload_dir="ignored",
            chunks=1,
            changed=1,
            skipped_unchanged=0,
            upserted=1,
            stale_deleted=0,
            stale_chunk_ids=[],
            namespace_vector_count=1,
            chunk_ids=["x::y::0"],
            changed_chunk_ids=["x::y::0"],
            source_files=["x.md"],
        )
        base.update(overrides)
        super().__init__(**base)


@pytest.fixture
def _client(tmp_path, monkeypatch):
    settings = _settings(tmp_path)
    monkeypatch.setattr(main, "get_settings", lambda: settings)
    fake_ingest = lambda settings=None: _FakeIngestSummary()
    monkeypatch.setattr(main, "ingest_corpus", fake_ingest)
    return TestClient(main.app), tmp_path


def _upload(client, filename, content: bytes, content_type="text/plain"):
    return client.post("/upload", files={"files": (filename, io.BytesIO(content), content_type)})


def test_upload_accepts_a_valid_markdown_file(_client, monkeypatch):
    client, tmp_path = _client
    called = {}

    def fake_ingest_corpus(settings=None):
        # NOTE: previously written as
        #   lambda settings=None: called.setdefault("ran", True) or _FakeIngestSummary(...)
        # which is broken: dict.setdefault(...) returns the value (True),
        # and `True or X` short-circuits to True without ever evaluating X
        # — so the mock returned the bare bool True instead of the fake
        # summary dict, and the endpoint's `summary["changed"]` blew up
        # with "'bool' object is not subscriptable". Use a real function
        # so the side effect and the return value don't fight each other.
        called["ran"] = True
        return _FakeIngestSummary(changed=1, upserted=1)

    monkeypatch.setattr(main, "ingest_corpus", fake_ingest_corpus)

    response = _upload(client, "notes.md", b"# Title\n\nSome content.\n")
    body = response.json()

    assert response.status_code == 200
    assert body["status"] == "success"
    assert body["files"] == ["notes.md"]
    assert body["rejected"] == []
    assert body["chunks_created"] == 1
    assert body["vectors_upserted"] == 1
    assert called.get("ran") is True
    assert (tmp_path / "notes.md").exists()


def test_upload_rejects_unsupported_extension(_client):
    client, tmp_path = _client
    response = _upload(client, "malware.exe", b"not really an exe", content_type="application/octet-stream")
    body = response.json()

    assert response.status_code == 200
    assert body["status"] == "error"
    assert body["files"] == []
    assert len(body["rejected"]) == 1
    assert body["rejected"][0]["filename"] == "malware.exe"
    assert "unsupported" in body["rejected"][0]["error"].lower()
    assert not (tmp_path / "malware.exe").exists()


def test_upload_rejects_empty_file(_client):
    client, tmp_path = _client
    response = _upload(client, "empty.txt", b"")
    body = response.json()

    assert body["status"] == "error"
    assert "empty" in body["rejected"][0]["error"].lower()
    assert not (tmp_path / "empty.txt").exists()


def test_upload_rejects_oversized_file(tmp_path, monkeypatch):
    # MAX_UPLOAD_FILE_SIZE_BYTES has a production-enforced minimum of 1024
    # bytes (see app/config.py) and must not be weakened for this test —
    # use the smallest valid limit, and a file that genuinely exceeds it.
    settings = _settings(tmp_path, MAX_UPLOAD_FILE_SIZE_BYTES=1024)
    monkeypatch.setattr(main, "get_settings", lambda: settings)
    monkeypatch.setattr(main, "ingest_corpus", lambda settings=None: _FakeIngestSummary())
    client = TestClient(main.app)

    oversized_content = b"x" * 2048  # 2048 bytes > the 1024-byte limit
    response = _upload(client, "big.txt", oversized_content)
    body = response.json()

    assert body["status"] == "error"
    assert "limit" in body["rejected"][0]["error"].lower()
    assert not (tmp_path / "big.txt").exists()


def test_upload_sanitizes_path_traversal_filename_and_stays_inside_upload_dir(_client):
    client, tmp_path = _client
    response = _upload(client, "../../.env", b"SECRET=should-not-land-here")
    body = response.json()

    assert response.status_code == 200
    # Either accepted under a sanitized name, or rejected outright — but it
    # must never be written as an actual ../../.env escaping tmp_path.
    assert not (tmp_path.parent.parent / ".env").exists()
    if body["files"]:
        saved_name = body["files"][0]
        assert saved_name != "../../.env"
        assert "/" not in saved_name and "\\" not in saved_name
        saved_path = tmp_path / saved_name
        assert saved_path.resolve().is_relative_to(tmp_path.resolve())


def test_upload_partial_failure_reports_both_accepted_and_rejected(_client):
    client, tmp_path = _client
    response = client.post(
        "/upload",
        files=[
            ("files", ("good.md", io.BytesIO(b"# Good\n\nContent.\n"), "text/markdown")),
            ("files", ("bad.exe", io.BytesIO(b"nope"), "application/octet-stream")),
        ],
    )
    body = response.json()

    assert response.status_code == 200
    assert body["status"] == "partial_success"
    assert body["files"] == ["good.md"]
    assert len(body["rejected"]) == 1
    assert body["rejected"][0]["filename"] == "bad.exe"
    assert (tmp_path / "good.md").exists()
    assert not (tmp_path / "bad.exe").exists()


def test_upload_corrupt_pdf_is_rejected_with_a_clean_error(_client):
    client, tmp_path = _client
    # Not a real PDF -> pypdf's PdfReader raises inside app.loaders.load_text,
    # which the endpoint catches as a LoaderError and reports cleanly.
    response = _upload(client, "corrupt.pdf", b"this is not a valid pdf file", content_type="application/pdf")
    body = response.json()

    assert response.status_code == 200
    assert body["status"] == "error"
    assert body["rejected"][0]["filename"] == "corrupt.pdf"
    assert not (tmp_path / "corrupt.pdf").exists()


def test_upload_response_never_breaks_ask_response_shape(_client):
    """The upload endpoint must be purely additive: it must not touch
    AskResponse's shape/behavior."""
    import inspect

    from app.schemas import AskResponse

    fields = set(AskResponse.model_fields)
    assert fields == {"answer", "found", "citations", "trace"}


def test_upload_with_no_files_is_a_client_error(_client):
    client, _ = _client
    response = client.post("/upload", files={})
    assert response.status_code in (400, 422)


def test_upload_accepts_a_txt_file(_client):
    client, tmp_path = _client
    response = _upload(client, "raw_notes.txt", b"Plain text notes, no markdown headings.")
    body = response.json()

    assert body["status"] == "success"
    assert body["files"] == ["raw_notes.txt"]
    assert (tmp_path / "raw_notes.txt").exists()


def test_upload_accepts_a_docx_file(_client, monkeypatch):
    """.docx extraction is monkeypatched the same way tests/test_loaders.py
    does it — a corrupt/garbage byte string is not a real DOCX (a zip
    archive), so python-docx would fail on it regardless of endpoint logic;
    this isolates the endpoint's own accept/reject/save behavior from
    python-docx's actual parser."""
    client, tmp_path = _client

    class _FakeParagraph:
        def __init__(self, text):
            self.text = text

    class _FakeDocxDocument:
        def __init__(self):
            self.paragraphs = [_FakeParagraph("Employment agreement excerpt."), _FakeParagraph("Notice period: 60 days.")]
            self.tables = []

    import docx

    monkeypatch.setattr(docx, "Document", lambda path_str: _FakeDocxDocument())

    response = _upload(client, "agreement.docx", b"PK\x03\x04fake-docx-bytes", content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    body = response.json()

    assert body["status"] == "success"
    assert body["files"] == ["agreement.docx"]
    assert (tmp_path / "agreement.docx").exists()


def test_upload_accepts_a_pdf_file(_client, monkeypatch):
    """Same isolation approach as the .docx test above, via pypdf.PdfReader."""
    client, tmp_path = _client

    class _FakePage:
        def extract_text(self):
            return "Lease clause: monthly rent is $2,000."

    class _FakePdfReader:
        def __init__(self, path_str):
            self.pages = [_FakePage()]

    import pypdf

    monkeypatch.setattr(pypdf, "PdfReader", _FakePdfReader)

    response = _upload(client, "lease.pdf", b"%PDF-fake-bytes", content_type="application/pdf")
    body = response.json()

    assert body["status"] == "success"
    assert body["files"] == ["lease.pdf"]
    assert (tmp_path / "lease.pdf").exists()


def test_upload_reingesting_the_same_unchanged_file_reports_zero_new_chunks(_client, monkeypatch):
    """Duplicate/unchanged upload: uploading the identical file content
    twice must not re-embed anything the second time — ingest_corpus (the
    real one, not a fixture-level fake) already guarantees this via
    content-hash skipping (see test_ingestion.py); this test only confirms
    the endpoint surfaces `unchanged_chunks` from whatever ingest_corpus
    reports, rather than always claiming everything as newly created."""
    client, tmp_path = _client
    monkeypatch.setattr(
        main,
        "ingest_corpus",
        lambda settings=None: _FakeIngestSummary(changed=0, upserted=0, skipped_unchanged=1),
    )

    response = _upload(client, "notes.md", b"# Notes\n\nSame content every time.\n")
    body = response.json()

    assert body["status"] == "success"
    assert body["chunks_created"] == 0
    assert body["unchanged_chunks"] == 1


# ---------------------------------------------------------------------
# Prompt_11.docx: uploads must land in data/uploads/, never data/corpus/,
# and must never contaminate the immutable six-document assignment corpus.
# ---------------------------------------------------------------------


def test_upload_never_writes_into_the_real_corpus_dir(_client, monkeypatch):
    """Requirement #2: runtime uploads must not contaminate data/corpus/.

    This deliberately leaves CORPUS_DIR at its default, so it resolves to
    the project's real `data/corpus/` — the same directory
    test_only_the_six_known_source_files_appear (test_chunking.py) checks.
    If /upload ever regresses to writing there again, this test fails
    immediately and the file is cleaned up either way."""
    client, upload_dir = _client
    from app.ingestion import _resolve_corpus_dir

    # Build the settings the same way `_settings()` above does (CORPUS_DIR
    # left at its default), without touching the real cached
    # `app.config.get_settings()` — this only needs the resolved path, not
    # a live singleton, and avoids any dependency on real API keys being
    # present in the environment.
    default_settings = _settings(tmp_path=upload_dir)
    real_corpus_dir = _resolve_corpus_dir(default_settings)
    contaminant = real_corpus_dir / "should_not_land_here.md"
    assert not contaminant.exists(), "test precondition: corpus dir must start clean"

    try:
        response = _upload(client, "should_not_land_here.md", b"# Test\n\nContent.\n")
        assert response.status_code == 200
        assert not contaminant.exists(), (
            "POST /upload must never write into the immutable data/corpus/ directory"
        )
        assert (upload_dir / "should_not_land_here.md").exists()
    finally:
        contaminant.unlink(missing_ok=True)


def test_upload_stores_files_under_the_configured_upload_dir(_client):
    """Requirement #1: /upload stores runtime documents under data/uploads/
    (here, the isolated tmp_path standing in for it via UPLOAD_DIR)."""
    client, upload_dir = _client
    response = _upload(client, "runtime_doc.md", b"# Runtime\n\nContent.\n")
    body = response.json()

    assert response.status_code == 200
    assert body["status"] == "success"
    assert (upload_dir / "runtime_doc.md").exists()


def test_upload_calls_ingest_corpus_without_a_directory_override(_client, monkeypatch):
    """The endpoint must rely on ingest_corpus()'s own default resolution
    of BOTH corpus_dir and upload_dir from settings — not pass a single
    directory override that would silently exclude the other knowledge
    base from reconciliation (see app/ingestion.py's unified-ingestion
    docstring)."""
    client, tmp_path = _client
    captured_kwargs = {}

    def fake_ingest_corpus(**kwargs):
        captured_kwargs.update(kwargs)
        return _FakeIngestSummary()

    monkeypatch.setattr(main, "ingest_corpus", fake_ingest_corpus)

    _upload(client, "notes.md", b"# Notes\n\nContent.\n")

    assert "corpus_dir" not in captured_kwargs
    assert "upload_dir" not in captured_kwargs
    assert "settings" in captured_kwargs